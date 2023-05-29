import datetime
from io import BytesIO

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches


class DocsGenerator:
    def __init__(self, static_data: dict):
        self.header = static_data["header"]
        self.footer = static_data["footer"]

    def generate_docs(self, patient_data: dict, gpt_sympthoms: str, gpt_recommendations: str):
        document = Document()

        sections = document.sections
        for section in sections:
            section.bottom_margin = Inches(0.5)

        logo_paragraph = document.add_paragraph()
        logo_run = logo_paragraph.add_run()
        logo_run.add_picture("./static/logo.png", width=Inches(1))

        title_paragraph = document.add_paragraph()
        title_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        title_run = title_paragraph.add_run("KARTA WIZYTY")
        title_run.bold = True

        patient = document.add_paragraph()
        patient.add_run("Imie: ").bold = True
        patient.add_run(patient_data["name"] + "\n")
        patient.add_run("Nazwisko: ").bold = True
        patient.add_run(patient_data["surname"] + "\n")
        patient.add_run("PESEL: ").bold = True
        patient.add_run(patient_data["pesel"] + "\n")
        patient.add_run("Adres: ").bold = True
        patient.add_run(patient_data["address"] + "\n")
        patient.add_run("Data wygenerowania dokumentu: ").bold = True
        patient.add_run(str(datetime.datetime.today()) + "\n")

        document.add_heading("WYWIAD Z PACJENTEM", level=1)
        document.add_paragraph(gpt_sympthoms)

        document.add_heading("ZALECENIA", level=1)
        document.add_paragraph(gpt_recommendations)

        footer_paragraph = document.sections[-1].footer.paragraphs[0]
        footer_paragraph.text = self.footer
        footer_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

        document_buffer = BytesIO()
        document.save(document_buffer)
        document_buffer.seek(0)
        return document_buffer
